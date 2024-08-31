from flask import Flask, request, jsonify, render_template, send_file
import speech_recognition as sr
from pydub import AudioSegment
from googletrans import Translator
import os
from fpdf import FPDF 
from docx import Document

app = Flask(__name__)
recognizer = sr.Recognizer()
translator = Translator()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/recognize', methods=['POST'])
def recognize():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'})

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'})

    language = request.form.get('language', 'en')
    audio_file_path = 'uploads/temp_audio.wav'
    
    audio = AudioSegment.from_file(file)
    audio.export(audio_file_path, format='wav')

    try:
        with sr.AudioFile(audio_file_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
            translated_text = translator.translate(text, dest=language).text
    except Exception as e:
        return jsonify({'error': str(e)})

    # Clean up the temporary file
    os.remove(audio_file_path)

    response_data = {
        'original_text': text,
        'translated_text': translated_text
    }

    return jsonify(response_data)

@app.route('/generate-file', methods=['POST'])
def generate_file():
    file_type = request.form.get('file_type')
    original_text = request.form.get('original_text')
    translated_text = request.form.get('translated_text')

    if not original_text or not translated_text:
        return jsonify({'error': 'No transcribed text available'})

    if file_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, f"Original Text:\n{original_text}\n\nTranslated Text:\n{translated_text}")
        pdf_file_path = 'outputs/transcription.pdf'
        pdf.output(pdf_file_path)
        return send_file(pdf_file_path, as_attachment=True)

    elif file_type == 'word':
        doc = Document()
        doc.add_heading('Transcription', 0)
        doc.add_paragraph(f"Original Text:\n{original_text}\n\nTranslated Text:\n{translated_text}")
        word_file_path = 'outputs/transcription.docx'
        doc.save(word_file_path)
        return send_file(word_file_path, as_attachment=True)

    elif file_type == 'txt':
        txt_file_path = 'outputs/transcription.txt'
        with open(txt_file_path, 'w') as f:
            f.write(f"Original Text:\n{original_text}\n\nTranslated Text:\n{translated_text}")
        return send_file(txt_file_path, as_attachment=True)

    return jsonify({'error': 'Invalid file type'})

if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    if not os.path.exists('outputs'):
        os.makedirs('outputs')
    app.run(debug=True)
